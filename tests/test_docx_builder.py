import pytest
import os
from core.docx_builder import DOCXBuilder
from docx import Document

def test_docx_builder_creation():
    builder = DOCXBuilder()
    builder.create_document()
    assert builder.doc is not None

def test_add_mixed_paragraph():
    builder = DOCXBuilder()
    builder.create_document()
    
    builder.add_mixed_paragraph("هذا اختبار", confidence=0.9)
    builder.add_mixed_paragraph("This is English", confidence=0.4)
    
    assert len(builder.doc.paragraphs) == 2
    assert "هذا اختبار" in builder.doc.paragraphs[0].text
    assert "This is English" in builder.doc.paragraphs[1].text

def test_add_scanned_blocks():
    builder = DOCXBuilder()
    builder.create_document()
    
    blocks = [
        {"type": "text", "text": "فقرة أولى"},
        {"type": "text", "text": "فقرة ثانية"}
    ]
    
    builder.add_scanned_blocks(blocks)
    assert len(builder.doc.paragraphs) == 2
    assert "فقرة ثانية" in builder.doc.paragraphs[1].text

def test_save_document(tmp_path):
    builder = DOCXBuilder()
    builder.create_document()
    builder.add_mixed_paragraph("حفظ المستند")
    
    out_file = tmp_path / "test.docx"
    builder.save(str(out_file))
    
    assert out_file.exists()
    
    # Verify content
    doc = Document(str(out_file))
    assert len(doc.paragraphs) == 1
    assert "حفظ المستند" in doc.paragraphs[0].text
