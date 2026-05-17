import pytest
from docx.oxml.ns import qn
from core.docx_builder import DOCXBuilder
from core.rtl_helpers import apply_rtl_run
from docx import Document

class TestStage3Assembly:
    """
    Stage 3 Tests: Verify that DOCX assembly creates correct and non-redundant XML,
    properly handles formatting (fonts, colors, styles), and applies correct RTL tags.
    """
    
    def test_no_duplicate_rfonts(self):
        """Test B3: Each run should have at most one w:rFonts element to prevent XML corruption."""
        builder = DOCXBuilder()
        builder.create_document()
        
        # Add a block that simulates what the scanner returns
        blocks = [{"type": "paragraph", "text": "اختبار", "spans": [
            {"text": "اختبار", "size": 12, "is_bold": False, "is_italic": False, "color": 0}
        ], "bbox": [0, 0, 100, 20]}]
        
        builder.add_scanned_blocks(blocks)
        
        # Verify the XML
        for p in builder.doc.paragraphs:
            for run in p.runs:
                rPr = run._element.find(qn('w:rPr'))
                if rPr is not None:
                    rfonts = rPr.findall(qn('w:rFonts'))
                    assert len(rfonts) <= 1, f"Found {len(rfonts)} w:rFonts elements! Should be at most 1."

    def test_color_correct_rgb(self):
        """Test B5: Colors from PyMuPDF (0xRRGGBB vs 0xBBGGRR) should be handled correctly."""
        builder = DOCXBuilder()
        builder.create_document()
        
        # PyMuPDF color 0xFF0000 means RED. (It represents 0xRRGGBB in fitz integer colors)
        blocks = [{"type": "paragraph", "text": "أحمر", "spans": [
            {"text": "أحمر", "size": 12, "is_bold": False, "is_italic": False, "color": 0xFF0000}
        ], "bbox": [0, 0, 100, 20]}]
        
        builder.add_scanned_blocks(blocks)
        
        run = builder.doc.paragraphs[0].runs[0]
        # It should be pure red
        assert run.font.color.rgb[0] == 0xFF, "Red channel should be FF"
        assert run.font.color.rgb[1] == 0x00, "Green channel should be 00"
        assert run.font.color.rgb[2] == 0x00, "Blue channel should be 00"

    def test_original_font_preserved(self):
        """Test B7: The font from the scanner should be preserved instead of overridden by default."""
        builder = DOCXBuilder(font_name="Arial") # Default font is Arial
        builder.create_document()
        
        # We specify 'Times New Roman' in the scanner output
        blocks = [{"type": "paragraph", "text": "اختبار", "spans": [
            {"text": "اختبار", "size": 12, "is_bold": False, "is_italic": False, "color": 0, "font": "Times New Roman"}
        ], "bbox": [0, 0, 100, 20]}]
        
        builder.add_scanned_blocks(blocks)
        
        run = builder.doc.paragraphs[0].runs[0]
        rPr = run._element.rPr
        rFonts = rPr.find(qn('w:rFonts'))
        
        # The Arabic font should be 'Times New Roman', not 'Arial'
        assert rFonts.get(qn('w:cs')) == "Times New Roman", f"Font was overridden. Got {rFonts.get(qn('w:cs'))}"
