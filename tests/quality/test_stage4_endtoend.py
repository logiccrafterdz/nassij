import pytest
from pathlib import Path
from docx import Document
from cli import convert_pdf_to_docx
from utils.metrics import calculate_cer

class TestStage4EndToEnd:
    """
    Stage 4 Tests: Complete pipeline tests (PDF -> DOCX).
    Ensures that the entire Nassij process produces high-quality output.
    """
    
    @pytest.mark.skip(reason="PyMuPDF generation without a TTF font replaces Arabic chars with '?'. E2E needs a real PDF corpus.")
    def test_e2e_simple_arabic(self, pdf_generator, tmp_path):
        """Test F1: Simple Arabic paragraph roundtrip."""
        source_text = "هذا نص عربي بسيط للاختبار. يجب أن يكون طول هذا النص أكثر من خمسين حرفاً لكي لا يعتبره النظام صفحة مصورة."
        pdf_path = pdf_generator.with_text(source_text, rtl=True)
        docx_path = tmp_path / "output.docx"
        
        # Run conversion in 'scan' mode
        success = convert_pdf_to_docx(str(pdf_path), str(docx_path), mode="scan")
        assert success, "Conversion failed"
        
        # Read resulting DOCX
        doc = Document(str(docx_path))
        extracted = " ".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        
        assert source_text in extracted
        
    @pytest.mark.skip(reason="PyMuPDF generation without a TTF font replaces Arabic chars with '?'. E2E needs a real PDF corpus.")
    def test_e2e_roundtrip_cer(self, pdf_generator, tmp_path):
        """Test F8: The golden test. CER should be extremely low for clean digital PDFs."""
        source_text = "بسم الله الرحمن الرحيم. الحمد لله رب العالمين. هذا النص طويل بما فيه الكفاية لكي لا يتم تصنيفه كصورة ضوئية."
        pdf_path = pdf_generator.with_text(source_text, rtl=True)
        docx_path = tmp_path / "golden.docx"
        
        success = convert_pdf_to_docx(str(pdf_path), str(docx_path), mode="scan")
        assert success
        
        doc = Document(str(docx_path))
        extracted = " ".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        
        cer = calculate_cer(source_text, extracted)
        # For a pure digital PDF with no complex layout, CER should be 0.
        # But we allow a small margin for space discrepancies.
        assert cer < 0.05, f"CER = {cer:.2%} is too high. Extracted: {extracted}"
