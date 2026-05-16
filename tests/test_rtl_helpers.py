import pytest
from docx import Document
from core.rtl_helpers import apply_rtl_paragraph, apply_rtl_run, apply_rtl_table
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def test_apply_rtl_paragraph():
    doc = Document()
    p = doc.add_paragraph()
    
    apply_rtl_paragraph(p, alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    
    assert p.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    pPr = p._element.pPr
    assert pPr is not None
    
    # Check if bidi tag exists
    bidi = pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bidi")
    assert bidi is not None
    assert bidi.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") in ("1", None)

def test_apply_rtl_run():
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("اختبار")
    
    apply_rtl_run(run, font_name="Arial", lang="ar-SA")
    
    assert run.font.name == "Arial"
    
    rPr = run._element.rPr
    assert rPr is not None
    
    # Check if rtl tag exists
    rtl = rPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rtl")
    assert rtl is not None

def test_apply_rtl_table():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    
    apply_rtl_table(table)
    
    tblPr = table._element.tblPr
    assert tblPr is not None
    
    # Check if bidiVisual exists
    bidiVisual = tblPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bidiVisual")
    assert bidiVisual is not None
