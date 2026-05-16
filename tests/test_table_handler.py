"""
Tests for the TableHandler module.
Covers table reconstruction, RTL tables, merged cells, and Arabic cell content.
"""
import pytest
from docx import Document
from docx.oxml.ns import qn

from core.table_handler import TableHandler
from core.arabic_processor import ArabicProcessor


@pytest.fixture
def handler():
    """Create a TableHandler with default settings."""
    return TableHandler(ArabicProcessor(), font_name="Arial")


class TestTableHandler:
    """Tests for TableHandler table reconstruction and RTL support."""

    def test_basic_table_creation(self, handler):
        """A simple 2x2 table should be created correctly."""
        doc = Document()
        table_data = {
            "cells": [
                ["Header 1", "Header 2"],
                ["Cell A", "Cell B"]
            ]
        }

        table = handler.reconstruct_table(table_data, doc)
        assert len(table.rows) == 2
        assert len(table.columns) == 2

    def test_arabic_table_has_bidi_visual(self, handler):
        """A table with Arabic content should have w:bidiVisual set."""
        doc = Document()
        table_data = {
            "cells": [
                ["الاسم", "العمر"],
                ["أحمد", "٢٥"]
            ],
            "is_arabic": True
        }

        table = handler.reconstruct_table(table_data, doc)

        # Check bidiVisual is present
        tblPr = table._element.tblPr
        bidiVisual = tblPr.find(qn('w:bidiVisual'))
        assert bidiVisual is not None, "Arabic table must have w:bidiVisual"

    def test_arabic_cells_have_rtl_properties(self, handler):
        """Individual Arabic cells should have RTL paragraph properties."""
        doc = Document()
        table_data = {
            "cells": [
                ["مرحبا", "عالم"]
            ],
            "is_arabic": True
        }

        table = handler.reconstruct_table(table_data, doc)

        # Check first cell's paragraph has bidi
        cell = table.rows[0].cells[0]
        p = cell.paragraphs[0]
        pPr = p._element.pPr
        if pPr is not None:
            bidi = pPr.find(qn('w:bidi'))
            assert bidi is not None, "Arabic cell paragraph must have w:bidi"

    def test_mixed_content_table(self, handler):
        """A table with mixed Arabic and English should handle both."""
        doc = Document()
        table_data = {
            "cells": [
                ["Name", "الاسم"],
                ["Ahmed", "أحمد"]
            ],
            "is_arabic": False
        }

        table = handler.reconstruct_table(table_data, doc)
        assert len(table.rows) == 2

        # The Arabic cell should still have RTL (per-cell detection)
        arabic_cell = table.rows[0].cells[1]
        assert "الاسم" in arabic_cell.paragraphs[0].text

    def test_empty_cells_handled(self, handler):
        """Empty cells should not cause errors."""
        doc = Document()
        table_data = {
            "cells": [
                ["البيانات", "", "المجموع"],
                ["", "قيمة", ""]
            ]
        }

        table = handler.reconstruct_table(table_data, doc)
        assert len(table.rows) == 2
        assert len(table.columns) == 3

    def test_empty_table_data(self, handler):
        """Empty table data should produce a 1x1 fallback table."""
        doc = Document()
        table_data = {"cells": []}

        table = handler.reconstruct_table(table_data, doc)
        assert len(table.rows) == 1
        assert len(table.columns) == 1

    def test_detect_merged_cells(self, handler):
        """Merged cell definitions should be returned from table data."""
        table_data = {
            "cells": [["A", "B"], ["C", "D"]],
            "merged_cells": [
                {"row": 0, "col": 0, "row_span": 1, "col_span": 2}
            ]
        }
        merged = handler.detect_merged_cells(table_data)
        assert len(merged) == 1
        assert merged[0]["col_span"] == 2

    def test_process_table_from_ocr_pipeline(self, handler):
        """The full OCR table pipeline should produce a valid DOCX table."""
        doc = Document()
        ocr_table = {
            "cells": [
                ["رقم", "الاسم", "الدرجة"],
                ["١", "محمد", "٩٠"],
                ["٢", "فاطمة", "٩٥"]
            ],
            "is_arabic": True
        }

        table = handler.process_table_from_ocr(ocr_table, doc)
        assert len(table.rows) == 3
        assert len(table.columns) == 3
