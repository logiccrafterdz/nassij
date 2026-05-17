"""
Nassij Pipeline Diagnostic Script
==================================
This script traces the ENTIRE conversion pipeline from PDF to DOCX,
dumping intermediate output at every stage to identify exactly where
text corruption happens.

Usage: python diagnose_pipeline.py <input.pdf>
"""
import sys
import json
from pathlib import Path

# Add project root
sys.path.insert(0, str(Path(__file__).parent))

from core.pdf_reader import PDFReader
from core.scanner import NassijScanner
from core.arabic_processor import ArabicProcessor
from core.docx_builder import DOCXBuilder
from docx import Document
from docx.oxml.ns import qn

def diagnose(pdf_path: str):
    print("=" * 80)
    print(f"  NASSIJ PIPELINE DIAGNOSTIC — {Path(pdf_path).name}")
    print("=" * 80)
    
    # ──────────────────────────────────────────────
    # STAGE 0: Raw PyMuPDF extraction (no Nassij code)
    # ──────────────────────────────────────────────
    print("\n" + "─" * 60)
    print("STAGE 0: Raw PyMuPDF get_text('text') — No Nassij code")
    print("─" * 60)
    
    import fitz
    doc = fitz.open(pdf_path)
    page = doc[0]
    
    raw_text = page.get_text("text")
    print(f"  Raw text length: {len(raw_text)} chars")
    print(f"  First 500 chars:\n  >{raw_text[:500]}<")
    
    # Check for mangled text indicators
    alif_sequences = raw_text.count("اا")
    print(f"\n  Double-Alif count: {alif_sequences}")
    print(f"  Contains '?' placeholders: {raw_text.count('?')}")
    
    # ──────────────────────────────────────────────
    # STAGE 0b: Raw PyMuPDF rawdict (character level)
    # ──────────────────────────────────────────────
    print("\n" + "─" * 60)
    print("STAGE 0b: Raw PyMuPDF get_text('rawdict') — Character level")
    print("─" * 60)
    
    rawdict = page.get_text("rawdict")
    blocks = rawdict.get("blocks", [])
    text_blocks = [b for b in blocks if b.get("type") == 0]
    
    print(f"  Total blocks: {len(blocks)}")
    print(f"  Text blocks: {len(text_blocks)}")
    
    if text_blocks:
        first_block = text_blocks[0]
        first_lines = first_block.get("lines", [])
        print(f"  First block has {len(first_lines)} lines")
        
        if first_lines:
            first_line = first_lines[0]
            spans = first_line.get("spans", [])
            print(f"  First line has {len(spans)} spans")
            
            for i, span in enumerate(spans[:3]):
                chars = span.get("chars", [])
                text_from_chars = "".join(c.get("c", "") for c in chars)
                font = span.get("font", "?")
                size = span.get("size", 0)
                print(f"  Span {i}: font='{font}', size={size:.1f}, text='{text_from_chars[:80]}'")
    
    # ──────────────────────────────────────────────
    # STAGE 1: NassijScanner output
    # ──────────────────────────────────────────────
    print("\n" + "─" * 60)
    print("STAGE 1: NassijScanner.scan_page() output")
    print("─" * 60)
    
    scanner = NassijScanner()
    scanned_blocks = scanner.scan_page(page)
    
    print(f"  Scanner produced {len(scanned_blocks)} blocks")
    
    for i, block in enumerate(scanned_blocks[:5]):
        btype = block.get("type", "?")
        btext = block.get("text", "")
        spans = block.get("spans", [])
        print(f"\n  Block {i} (type={btype}, spans={len(spans)}):")
        print(f"    Text: '{btext[:200]}'")
        
        if spans:
            for j, s in enumerate(spans[:3]):
                print(f"    Span {j}: font='{s.get('font', '')}', size={s.get('size', 0):.1f}, text='{s.get('text', '')[:60]}'")
    
    # ──────────────────────────────────────────────
    # STAGE 2: ArabicProcessor output
    # ──────────────────────────────────────────────
    print("\n" + "─" * 60)
    print("STAGE 2: ArabicProcessor.process_paragraph() output")
    print("─" * 60)
    
    processor = ArabicProcessor()
    
    for i, block in enumerate(scanned_blocks[:3]):
        if block.get("type") == "table":
            continue
        text = block.get("text", "")
        result = processor.process_paragraph(text, logical_output=True)
        print(f"\n  Block {i}:")
        print(f"    Input:  '{text[:200]}'")
        print(f"    Output: '{result['text'][:200]}'")
        print(f"    is_arabic: {result['is_arabic']}")
        
        # Check if text changed
        if text != result['text']:
            print(f"    ⚠️ TEXT CHANGED during processing!")
            # Show differences
            for ci, (a, b) in enumerate(zip(text, result['text'])):
                if a != b:
                    print(f"    Diff at char {ci}: '{a}' (U+{ord(a):04X}) → '{b}' (U+{ord(b):04X})")
                    if ci > 10:
                        print(f"    ... (more diffs)")
                        break
    
    # ──────────────────────────────────────────────
    # STAGE 3: DOCX Builder output
    # ──────────────────────────────────────────────
    print("\n" + "─" * 60)
    print("STAGE 3: DOCXBuilder.add_scanned_blocks() output")
    print("─" * 60)
    
    builder = DOCXBuilder()
    builder.create_document()
    builder.add_scanned_blocks(scanned_blocks)
    
    out_path = Path(pdf_path).with_suffix(".diagnostic.docx")
    builder.save(str(out_path))
    
    # Read back the DOCX to verify
    doc_out = Document(str(out_path))
    paragraphs = [p.text.strip() for p in doc_out.paragraphs if p.text.strip()]
    
    print(f"  DOCX has {len(doc_out.paragraphs)} paragraphs ({len(paragraphs)} non-empty)")
    
    for i, p in enumerate(doc_out.paragraphs[:5]):
        if not p.text.strip():
            continue
        print(f"\n  Paragraph {i}:")
        print(f"    Text: '{p.text[:200]}'")
        
        # Check XML properties
        pPr = p._element.pPr
        has_bidi = pPr is not None and pPr.find(qn('w:bidi')) is not None
        print(f"    Has w:bidi: {has_bidi}")
        
        for j, run in enumerate(p.runs[:3]):
            rPr = run._element.rPr
            has_rtl = rPr is not None and rPr.find(qn('w:rtl')) is not None
            rfonts_count = len(rPr.findall(qn('w:rFonts'))) if rPr is not None else 0
            rfonts_cs = None
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    rfonts_cs = rf.get(qn('w:cs'))
            print(f"    Run {j}: rtl={has_rtl}, rFonts_count={rfonts_count}, cs_font='{rfonts_cs}', text='{run.text[:60]}'")
    
    # ──────────────────────────────────────────────
    # SUMMARY
    # ──────────────────────────────────────────────
    print("\n" + "=" * 80)
    print("  DIAGNOSTIC SUMMARY")
    print("=" * 80)
    
    # Compare raw vs final
    raw_clean = raw_text.replace("\n", " ").strip()
    final_clean = " ".join(paragraphs)
    
    from utils.metrics import calculate_cer
    cer = calculate_cer(raw_clean, final_clean)
    
    print(f"  Raw text length:   {len(raw_clean)}")
    print(f"  Final DOCX length: {len(final_clean)}")
    print(f"  CER (Raw vs DOCX): {cer:.2%}")
    print(f"  Output saved to:   {out_path}")
    
    doc.close()
    
    print("\n  Done.")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python diagnose_pipeline.py <input.pdf>")
        sys.exit(1)
    diagnose(sys.argv[1])
