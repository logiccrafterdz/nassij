"""
Command-line interface for Nassij PDF-to-DOCX converter.
"""
import argparse
import sys
from pathlib import Path
from typing import Optional, Callable
import time

import logging
from core.pdf_reader import PDFReader
from core.ocr_engine import OCRFacade
from core.docx_builder import DOCXBuilder
from core.scanner import NassijScanner
from utils.metrics import calculate_all_metrics
from integrity.proof import IntegrityProof
import json

logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
logger = logging.getLogger(__name__)

def convert_pdf_to_docx(
    input_pdf: str,
    output_docx: str,
    mode: str = 'balanced',
    preserve_diacritics: bool = True,
    font_name: str = 'Arial',
    dpi: int = 300,
    generate_proof: bool = False,
    progress_callback: Optional[Callable[[int, int, str], None]] = None
) -> bool:
    """
    Convert PDF to DOCX with Arabic support.
    
    Args:
        input_pdf: Path to input PDF file
        output_docx: Path to output DOCX file
        mode: Conversion mode ('fast', 'balanced', 'accurate')
        preserve_diacritics: Whether to preserve Arabic diacritics
        font_name: Font name for Arabic text
        dpi: DPI for scanned page conversion
        
    Returns:
        True if conversion successful, False otherwise
    """
    try:
        logger.info(f"Reading PDF: {input_pdf}")
        
        # Initialize PDF reader
        with PDFReader(input_pdf) as pdf_reader:
            page_count = pdf_reader.get_page_count()
            logger.info(f"Found {page_count} page(s)")
            
            # Enforce mode settings
            if mode in ['legal', 'research']:
                generate_proof = True
                preserve_diacritics = True
                if mode == 'legal':
                    dpi = max(dpi, 400)  # Force higher DPI for legal OCR
                    
            # Initialize DOCX builder
            docx_builder = DOCXBuilder(
                font_name=font_name,
                preserve_diacritics=preserve_diacritics
            )
            docx_builder.create_document()
            
            # Initialize OCR engine or Scanner
            ocr_engine: Optional[OCRFacade] = None
            scanner: Optional[NassijScanner] = None
            proof = IntegrityProof() if generate_proof else None
            
            if mode in ('scan', 'research'):
                scanner = NassijScanner()
                logger.info(f"NassijScanner initialized (Mode: {mode})")
            elif mode in ('balanced', 'accurate', 'legal'):
                try:
                    ocr_engine = OCRFacade(lang='ar', use_table=True)
                    logger.info("OCR engine initialized")
                except Exception as e:
                    logger.warning(f"OCR engine initialization failed: {e}")
                    logger.warning("Falling back to text extraction only")
                    mode = 'fast'
            
            # Process each page
            for page_num in range(page_count):
                if progress_callback:
                    progress_callback(page_num + 1, page_count, "extracting")
                else:
                    logger.info(f"Processing page {page_num + 1}/{page_count}...")
                
                if mode == 'scan':
                    page = pdf_reader.doc[page_num]
                    page_text = page.get_text().strip()
                    
                    if len(page_text) < 50:
                        # Auto-fallback: this page is scanned, try OCR if available
                        logger.warning(f"  Page {page_num + 1} appears scanned. Falling back to OCR...")
                        if ocr_engine is None:
                            try:
                                ocr_engine = OCRFacade(lang='ar', use_table=True)
                            except Exception:
                                logger.warning(f"  OCR engine unavailable. Skipping scanned page {page_num + 1}.")
                                continue
                        if ocr_engine and ocr_engine.engine:
                            page_image = pdf_reader.convert_page_to_image(page_num, dpi=dpi)
                            ocr_result = ocr_engine.extract_from_pil_image(page_image)
                            if ocr_result['text_blocks']:
                                docx_builder.add_text_blocks(ocr_result['text_blocks'])
                                if proof:
                                    for b in ocr_result['text_blocks']:
                                        proof.add_block(b.get('text', ''), b.get('type', 'text'))
                        continue
                    
                    assert scanner is not None  # Guaranteed by mode == 'scan' branch above
                    blocks = scanner.scan_page(page)
                    if blocks:
                        docx_builder.add_scanned_blocks(blocks)
                        if proof:
                            for b in blocks:
                                proof.add_block(b.get('text', ''), b.get('type', 'text'))
                    continue
                
                # Extract text from page (legacy modes)
                page_data = pdf_reader.extract_text_from_page(page_num)
                
                # Check if page is scanned OR if accurate mode forces OCR
                # Accurate/Legal mode forces OCR to bypass potentially corrupted/visual PDF text extraction
                should_use_ocr = (mode in ['accurate', 'legal']) or (page_data['is_scanned'] and mode not in ['fast', 'scan', 'research'])
                
                if should_use_ocr:
                    if ocr_engine:
                        reason = "scanned page" if page_data['is_scanned'] else "accurate mode forced"
                        logger.info(f"  Page {page_num + 1} processing using OCR ({reason})...")
                        # Convert page to image
                        page_image = pdf_reader.convert_page_to_image(page_num, dpi=dpi)
                        
                        # Run OCR
                        ocr_result = ocr_engine.extract_from_pil_image(page_image)
                        
                        # Add text blocks
                        if ocr_result['text_blocks']:
                            docx_builder.add_text_blocks(ocr_result['text_blocks'])
                            if proof:
                                for b in ocr_result['text_blocks']:
                                    proof.add_block(b.get('text', ''), b.get('type', 'text'))
                        
                        # Add tables
                        for table_data in ocr_result['tables']:
                            docx_builder.add_table(table_data)
                    else:
                        logger.warning(f"  OCR requested but engine unavailable. Falling back to text.")
                        # Text-based fallback
                        if page_data['text'].strip():
                             if page_data['blocks']:
                                docx_builder.add_text_blocks(page_data['blocks'])
                                if proof:
                                    for b in page_data['blocks']:
                                        proof.add_block(b.get('text', ''), b.get('type', 'text'))
                             else:
                                docx_builder.add_mixed_paragraph(page_data['text'])
                                if proof:
                                    proof.add_block(page_data['text'], 'text')
                else:
                    # Text-based page
                    if page_data['text'].strip():
                        # Add text blocks
                        if page_data['blocks']:
                            docx_builder.add_text_blocks(page_data['blocks'])
                            if proof:
                                for b in page_data['blocks']:
                                    proof.add_block(b.get('text', ''), b.get('type', 'text'))
                        else:
                            # Fallback: add full text as paragraph
                            docx_builder.add_mixed_paragraph(page_data['text'])
                            if proof:
                                proof.add_block(page_data['text'], 'text')
            
            # Save document
            logger.info(f"Saving DOCX: {output_docx}")
            docx_builder.save(output_docx)
            
            if proof:
                proof_path = output_docx + ".nassij-proof"
                logger.info(f"Generating Linguistic Proof: {proof_path}")
                proof.generate_proof_file(Path(input_pdf).name, proof_path)
                
            logger.info("Conversion completed successfully!")
            return True
            
    except Exception as e:
        logger.error(f"Error during conversion: {e}", exc_info=True)
        return False


def main():
    """Main CLI entry point."""
    parser = argparse.ArgumentParser(
        description='Nassij: High-accuracy PDF-to-DOCX converter for Arabic content',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  nassij convert input.pdf -o output.docx
  nassij convert input.pdf -o output.docx --mode accurate --font "Noto Sans Arabic"
  nassij convert input.pdf -o output.docx --preserve-diacritics --dpi 400
        """
    )
    
    subparsers = parser.add_subparsers(dest='command', help='Command to execute')
    
    # Convert command
    convert_parser = subparsers.add_parser('convert', help='Convert PDF to DOCX')
    convert_parser.add_argument('input', type=str, help='Input PDF file path')
    convert_parser.add_argument('-o', '--output', type=str, required=True,
                               help='Output DOCX file path')
    convert_parser.add_argument('--mode', type=str, 
                               choices=['scan', 'fast', 'balanced', 'accurate', 'legal', 'research'],
                               default='scan',
                               help='Conversion mode (default: scan)')
    convert_parser.add_argument('--preserve-diacritics', action='store_true',
                               default=True,
                               help='Preserve Arabic diacritics (tashkeel)')
    convert_parser.add_argument('--no-preserve-diacritics', 
                               dest='preserve_diacritics',
                               action='store_false',
                               help='Do not preserve Arabic diacritics')
    convert_parser.add_argument('--font', type=str, default='Arial',
                               help='Font name for Arabic text (default: Arial)')
    convert_parser.add_argument('--dpi', type=int, default=300,
                               help='DPI for scanned page conversion (default: 300)')
    convert_parser.add_argument('--proof', action='store_true',
                               help='Generate a Linguistic Merkle Tree proof file')
    
    # Verify command
    verify_parser = subparsers.add_parser('verify', help='Verify Linguistic Proof')
    verify_parser.add_argument('docx', type=str, help='Output DOCX file to verify (for context)')
    verify_parser.add_argument('--proof', type=str, required=True,
                               help='Path to the .nassij-proof JSON file')
                               
    # Batch command
    batch_parser = subparsers.add_parser('batch', help='Batch convert multiple PDFs')
    batch_parser.add_argument('input_dir', type=str, help='Input directory containing PDF files')
    batch_parser.add_argument('-o', '--output_dir', type=str, required=True,
                               help='Output directory for DOCX files')
    batch_parser.add_argument('--mode', type=str, 
                               choices=['scan', 'fast', 'balanced', 'accurate', 'legal', 'research'],
                               default='scan')
    batch_parser.add_argument('--workers', type=int, default=4,
                               help='Number of parallel workers (default: 4)')
    batch_parser.add_argument('--proof', action='store_true',
                               help='Generate Linguistic Merkle Tree proof files')
    
    parser.add_argument('--info', type=str, metavar='FILE', help='Print PDF info without converting')
    parser.add_argument('--benchmark', type=str, metavar='FILE', help='Run a benchmark test on the given PDF')
    
    args = parser.parse_args()
    
    if args.info:
        print(f"Analyzing PDF: {args.info}")
        try:
            with PDFReader(args.info) as pdf_reader:
                pages = pdf_reader.get_page_count()
                print(f"Total Pages: {pages}")
                # Analyze a sample of pages to determine type
                scanned_pages = 0
                sample_size = min(pages, 5)
                for i in range(sample_size):
                    data = pdf_reader.extract_text_from_page(i)
                    if data['is_scanned']: scanned_pages += 1
                
                if scanned_pages == sample_size:
                    print("Document Type: Scanned (Requires OCR)")
                elif scanned_pages > 0:
                    print("Document Type: Hybrid (Contains scanned pages)")
                else:
                    print("Document Type: Digital Native")
        except Exception as e:
            print(f"Error reading PDF: {e}")
        return
        
    if args.benchmark:
        print(f"Running benchmark on: {args.benchmark}")
        import tempfile
        benchmark_output = Path(tempfile.mkdtemp()) / "benchmark_output.docx"
        start_time = time.time()
        success = convert_pdf_to_docx(
            input_pdf=args.benchmark, 
            output_docx=str(benchmark_output), 
            mode='scan', 
            preserve_diacritics=True, 
            font_name='Arial', 
            dpi=300
        )
        elapsed = time.time() - start_time
        if success:
            print(f"\nBenchmark completed successfully in {elapsed:.2f} seconds.")
            # Cleanup temp file
            benchmark_output.unlink(missing_ok=True)
        else:
            print(f"\nBenchmark failed after {elapsed:.2f} seconds.")
        return
    
    if not args.command:
        parser.print_help()
        sys.exit(1)
        
    if args.command == 'verify':
        print(f"Verifying Proof File: {args.proof} against {args.docx}")
        try:
            with open(args.proof, 'r', encoding='utf-8') as f:
                proof_data = json.load(f)
            
            # Recalculate root hash from leaves to verify tree integrity
            from integrity.merkle_tree import MerkleTree
            from integrity.diacritics_splitter import split_diacritics
            
            tree = MerkleTree()
            for block in proof_data.get('blocks', []):
                tree.add_leaf(block['compound_hash'])
            calculated_root = tree.build()
            
            if calculated_root != proof_data.get('merkle_root'):
                print("❌ Proof Tree is Invalid or Corrupted!")
                print(f"   Expected Root: {proof_data.get('merkle_root')}")
                print(f"   Calculated:    {calculated_root}")
                sys.exit(1)
                
            print("✅ Internal Proof Tree is Valid.")
            
            # Verify against DOCX
            if not Path(args.docx).exists():
                print(f"Error: DOCX file not found: {args.docx}")
                sys.exit(1)
                
            from docx import Document
            doc = Document(args.docx)
            
            # We only extract paragraphs because tables currently aren't hashed in the pipeline
            docx_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            
            expected_blocks = [b for b in proof_data.get('blocks', []) if b.get('type') == 'text']
            
            # Note: There might be slight count mismatches if empty blocks were filtered out differently.
            # A true robust verifier would do LCS matching, but for now we do direct matching.
            matched = 0
            for i, p_text in enumerate(docx_paragraphs):
                # We don't know the exact block_index if some were skipped, but let's try to find a match
                # by computing hash without block_index or by checking against all hashes.
                # Since we added block_index, let's just find if the base text matches.
                hashes = split_diacritics(p_text, block_index=-1) # calculate independent of index to see if text exists
                
                # We check if the base hash exists in the proof
                found = any(b['base_hash'] == hashes['base_hash'] for b in expected_blocks)
                if found:
                    matched += 1
                    
            match_percentage = (matched / len(expected_blocks)) * 100 if expected_blocks else 100
            
            print(f"   DOCX Match: {matched}/{len(expected_blocks)} blocks ({match_percentage:.1f}%)")
            
            if match_percentage >= 95.0:
                print("✅ DOCX Content Verified!")
                sys.exit(0)
            else:
                print("❌ DOCX Content Does Not Match Proof!")
                sys.exit(1)
                
        except Exception as e:
            print(f"Error reading proof file: {e}")
            sys.exit(1)
    
    if args.command == 'convert':
        # Validate input file
        input_path = Path(args.input)
        if not input_path.exists():
            print(f"Error: Input file not found: {args.input}", file=sys.stderr)
            sys.exit(1)
        
        # Validate output directory
        output_path = Path(args.output)
        output_dir = output_path.parent
        if output_dir and not output_dir.exists():
            output_dir.mkdir(parents=True, exist_ok=True)
        
        def cli_progress(current, total, phase):
            percent = int((current / total) * 100) if total > 0 else 100
            bar_len = 40
            filled = int(bar_len * current / total) if total > 0 else bar_len
            bar = '█' * filled + '-' * (bar_len - filled)
            sys.stdout.write(f'\r\033[K[{bar}] {current}/{total} {phase} ({percent}%)')
            sys.stdout.flush()
            if current == total:
                print()

        # Perform conversion
        start_time = time.time()
        success = convert_pdf_to_docx(
            input_pdf=str(input_path),
            output_docx=str(output_path),
            mode=args.mode,
            preserve_diacritics=args.preserve_diacritics,
            font_name=args.font,
            dpi=args.dpi,
            generate_proof=args.proof,
            progress_callback=cli_progress
        )
        elapsed = time.time() - start_time
        
        if success:
            print("\n" + "="*40)
            print("         Conversion Quality Card")
            print("="*40)
            print(f" Mode:       {args.mode}")
            print(f" Time:       {elapsed:.2f}s")
            print(f" Diacritics: {'Preserved' if args.preserve_diacritics else 'Removed'}")
            print(f" Proof:      {'Generated' if args.proof else 'Skipped'}")
            print("="*40)
        
        sys.exit(0 if success else 1)

    if args.command == 'batch':
        input_dir = Path(args.input_dir)
        output_dir = Path(args.output_dir)
        
        if not input_dir.exists() or not input_dir.is_dir():
            print(f"Error: Input directory not found: {args.input_dir}", file=sys.stderr)
            sys.exit(1)
            
        output_dir.mkdir(parents=True, exist_ok=True)
        pdf_files = list(input_dir.glob("*.pdf"))
        
        if not pdf_files:
            print(f"No PDF files found in {args.input_dir}")
            sys.exit(0)
            
        print(f"Found {len(pdf_files)} PDF files. Starting batch processing with {args.workers} workers...")
        
        import concurrent.futures
        
        def process_single(pdf_path):
            out_path = output_dir / (pdf_path.stem + ".docx")
            try:
                # Basic progress not shown per file in batch mode to avoid terminal spam
                success = convert_pdf_to_docx(
                    input_pdf=str(pdf_path),
                    output_docx=str(out_path),
                    mode=args.mode,
                    generate_proof=args.proof
                )
                return pdf_path.name, success
            except Exception as e:
                logger.error(f"Failed to process {pdf_path.name}: {e}")
                return pdf_path.name, False
                
        start_time = time.time()
        success_count = 0
        
        with concurrent.futures.ProcessPoolExecutor(max_workers=args.workers) as executor:
            futures = [executor.submit(process_single, pdf) for pdf in pdf_files]
            for future in concurrent.futures.as_completed(futures):
                name, success = future.result()
                status = "✅" if success else "❌"
                print(f"[{status}] {name}")
                if success: success_count += 1
                
        elapsed = time.time() - start_time
        print(f"\nBatch processing completed in {elapsed:.2f}s")
        print(f"Successful: {success_count}/{len(pdf_files)}")
        sys.exit(0 if success_count == len(pdf_files) else 1)
if __name__ == '__main__':
    main()

