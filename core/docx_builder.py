"""
DOCX generation module with full RTL support.
Creates Microsoft Word-compatible documents with proper Arabic rendering.
"""
from docx import Document
from docx.document import Document as DocumentType
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import Optional, List, Dict

from core.arabic_processor import ArabicProcessor
from core.table_handler import TableHandler
from core.rtl_helpers import apply_rtl_paragraph, apply_rtl_run


class DOCXBuilder:
    """
    Builds RTL-compliant DOCX files with Arabic text support.
    """
    
    def __init__(self, 
                 font_name: str = "Arial",
                 preserve_diacritics: bool = True):
        """
        Initialize DOCX builder.
        
        Args:
            font_name: Font name for Arabic text (default: Arial)
            preserve_diacritics: Whether to preserve diacritics
        """
        self.font_name = font_name
        self.arabic_processor = ArabicProcessor(preserve_diacritics=preserve_diacritics)
        self.table_handler = TableHandler(self.arabic_processor, font_name=font_name)
        self.doc: Optional[DocumentType] = None
    
    def create_document(self) -> DocumentType:
        """
        Create a new DOCX document and set section-level RTL.
        """
        self.doc = Document()
        
        # Set section RTL (Crucial for page margins and general direction)
        for section in self.doc.sections:
            sectPr = section._sectPr
            bidi = sectPr.find(qn('w:bidi'))
            if bidi is None:
                bidi = OxmlElement('w:bidi')
                sectPr.append(bidi)
                
        return self.doc
    
    def add_rtl_paragraph(self, 
                         text: str, 
                         font_name: Optional[str] = None,
                         font_size: int = 11) -> None:
        """
        Add a right-to-left paragraph with Arabic text support.
        
        Args:
            text: Text content
            font_name: Font name (defaults to instance font_name)
            font_size: Font size in points
        """
        if not self.doc:
            raise RuntimeError("Document not created. Call create_document() first.")
        
        font_name = font_name or self.font_name
        
        # Process Arabic text in LOGICAL order for DOCX
        # Word handles shaping and bidi reordering itself
        processed = self.arabic_processor.process_paragraph(text, logical_output=True)
        processed_text = processed['text']
        is_arabic = processed['is_arabic']
        
        # Create paragraph
        p = self.doc.add_paragraph()
        
        # Set RTL if Arabic
        if is_arabic:
            apply_rtl_paragraph(p)
        else:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Add text with font settings
        run = p.add_run(processed_text)
        run.font.size = Pt(font_size)
        
        # Set RTL run property for Arabic
        if is_arabic:
            apply_rtl_run(run, font_name)
        else:
            run.font.name = font_name
    
    def add_mixed_paragraph(self, 
                           text: str,
                           font_name: Optional[str] = None,
                           font_size: int = 11,
                           confidence: float = 1.0) -> None:
        """
        Add a paragraph that may contain mixed Arabic and Latin text.
        Automatically detects script and applies appropriate formatting.
        
        Args:
            text: Text content (may be mixed)
            font_name: Font name
            font_size: Font size in points
            confidence: OCR confidence score (0.0 to 1.0)
        """
        if not self.doc:
            raise RuntimeError("Document not created. Call create_document() first.")
        
        font_name = font_name or self.font_name
        
        # Process text in LOGICAL order for Word
        processed = self.arabic_processor.process_paragraph(text, logical_output=True)
        processed_text = processed['text']
        is_arabic = processed['is_arabic']
        
        # Create paragraph
        p = self.doc.add_paragraph()
        
        # For mixed text, use RTL if primarily Arabic
        if is_arabic:
            apply_rtl_paragraph(p)
            run = p.add_run(processed_text)
            run.font.size = Pt(font_size)
            apply_rtl_run(run, font_name)
        else:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = p.add_run(processed_text)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            
        # Highlight low confidence text
        if confidence < 0.5:
            run.font.highlight_color = WD_COLOR_INDEX.RED
        elif confidence < 0.85:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    
    def add_image(self, image_data: bytes, width_points: Optional[float] = None) -> None:
        """
        Add an image to the document.
        """
        from io import BytesIO
        from docx.shared import Inches
        
        if not self.doc:
            raise RuntimeError("Document not created.")
            
        # Add a paragraph for the image
        p = self.doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        run = p.add_run()
        # Word DPI is 72, so points / 72 = inches
        width = Inches(width_points / 72.0) if width_points else None
        run.add_picture(BytesIO(image_data), width=width)

    def add_table(self, table_data: Dict) -> None:
        """
        Add a table to the document.
        """
        if not self.doc:
            raise RuntimeError("Document not created.")
        self.table_handler.process_table_from_ocr(table_data, self.doc)

    def add_text_blocks(self, text_blocks: List[Dict]) -> None:
        """
        Add multiple blocks (text, table, or image) to the document.
        """
        if not self.doc:
            raise RuntimeError("Document not created.")
        
        for block in text_blocks:
            b_type = block.get('type', 'text')
            if b_type == 'text':
                text = block.get('text', '').strip()
                confidence = block.get('confidence', 1.0)
                if text:
                    # Paragraph direction
                    self.add_mixed_paragraph(text, confidence=confidence)
            elif b_type == 'table':
                self.add_table(block)
            elif b_type == 'image':
                img_bytes = block.get('image_bytes')
                if img_bytes:
                    bbox = block.get('bbox', [0,0,100,0])
                    width = bbox[2] - bbox[0]
                    self.add_image(img_bytes, width_points=width)

    def add_scanned_blocks(self, blocks: List[Dict]) -> None:
        """
        Add rich blocks extracted by NassijScanner.
        Preserves font size, bold, italic, and color per run.
        """
        if not self.doc:
            raise RuntimeError("Document not created.")
            
        for block in blocks:
            # Route table blocks to the table handler
            if block.get('type') == 'table':
                self.add_table(block)
                continue
            
            # Create paragraph
            p = self.doc.add_paragraph()
            
            # Determine if Arabic based on full text
            full_text = block.get('text', '')
            processed = self.arabic_processor.process_paragraph(full_text, logical_output=True)
            is_arabic = processed['is_arabic']
            
            # Paragraph level RTL
            if is_arabic:
                apply_rtl_paragraph(p)
            else:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                
            for span in block.get('spans', []):
                span_text = span.get('text', '')
                if not span_text.strip():
                    continue
                    
                processed_span = self.arabic_processor.process_paragraph(span_text, logical_output=True)
                run = p.add_run(processed_span['text'])
                
                # Apply styles
                run.font.size = Pt(span.get('size', 11))
                run.bold = span.get('is_bold', False)
                run.italic = span.get('is_italic', False)
                
                # Color
                color_int = span.get('color', 0)
                r = (color_int >> 16) & 255
                g = (color_int >> 8) & 255
                b = color_int & 255
                run.font.color.rgb = RGBColor(r, g, b)
                
                span_font = span.get('font', '').strip()
                target_font = span_font if span_font else self.font_name
                
                if is_arabic:
                    apply_rtl_run(run, target_font)
                else:
                    run.font.name = target_font
    
    def save(self, output_path: str) -> str:
        """
        Save document to file.
        
        Args:
            output_path: Path to save DOCX file
            
        Returns:
            Path to saved file
        """
        if not self.doc:
            raise RuntimeError("Document not created. Call create_document() first.")
        
        self.doc.save(output_path)
        return output_path

