"""
Table reconstruction module.
Rebuilds tables from OCR cell coordinates with Arabic text support.
"""
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import List, Dict, Optional, Tuple
from docx import Document
from docx.table import Table, _Cell
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.arabic_processor import ArabicProcessor
from utils.unicode_helpers import is_arabic_char
from core.rtl_helpers import apply_rtl_paragraph, apply_rtl_run, apply_rtl_table


class TableHandler:
    """
    Handles table reconstruction from OCR results.
    Supports merged cells, mixed Arabic/English, and borderless tables.
    """
    
    def __init__(self, arabic_processor: Optional[ArabicProcessor] = None, font_name: str = 'Arial'):
        """
        Initialize table handler.
        
        Args:
            arabic_processor: ArabicProcessor instance for text normalization
            font_name: Font name for cell text rendering
        """
        self.arabic_processor = arabic_processor or ArabicProcessor()
        self.font_name = font_name
    
    def reconstruct_table(self, table_data: Dict, doc: Document) -> Table:
        """
        Reconstruct a table in DOCX from OCR table data.
        """
        cells = table_data.get('cells', [])
        if not cells:
            return doc.add_table(rows=1, cols=1)
        
        num_rows = len(cells)
        max_cols = max(len(row) for row in cells) if cells else 1
        
        table = doc.add_table(rows=num_rows, cols=max_cols)
        # Use a more standard style that definitely has borders
        table.style = 'Table Grid'
        
        # Check if table is primarily Arabic
        is_arabic_table = table_data.get('is_arabic', False)
        if not is_arabic_table:
            # Fallback heuristic if flag not present
            arabic_cells = 0
            total_cells = 0
            for row in cells:
                for cell in row:
                    total_cells += 1
                    if any(is_arabic_char(c) for c in str(cell)):
                        arabic_cells += 1
            is_arabic_table = (arabic_cells / total_cells > 0.3) if total_cells > 0 else False

        # Add bidiVisual for RTL table
        if is_arabic_table:
            apply_rtl_table(table)
        
        for row_idx, row_data in enumerate(cells):
            for col_idx, cell_text in enumerate(row_data):
                if row_idx < num_rows and col_idx < max_cols:
                    cell = table.rows[row_idx].cells[col_idx]
                    
                    if not cell_text:
                        continue
                        
                    # Process Arabic text
                    processed = self.arabic_processor.process_paragraph(str(cell_text), logical_output=True)
                    cell_text_processed = processed['text']
                    
                    # Clear default paragraph
                    if cell.paragraphs:
                        p = cell.paragraphs[0]
                    else:
                        p = cell.add_paragraph()
                    
                    # Alignment: For Arabic tables, default cells to RIGHT
                    if is_arabic_table or processed['is_arabic']:
                        apply_rtl_paragraph(p, WD_ALIGN_PARAGRAPH.RIGHT)
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    run = p.add_run(cell_text_processed)
                    run.font.size = Pt(10)
                    
                    if is_arabic_table or processed['is_arabic']:
                        apply_rtl_run(run, self.font_name)
                    else:
                        run.font.name = self.font_name
        
        return table
    
    def detect_merged_cells(self, table_data: Dict) -> List[Dict]:
        """
        Detect merged cells from OCR bounding boxes.
        This is a heuristic approach - actual implementation depends on OCR output.
        
        Args:
            table_data: Table data with cell bounding boxes
            
        Returns:
            List of merged cell definitions:
            [{'row': int, 'col': int, 'row_span': int, 'col_span': int}, ...]
        """
        # If the OCR engine already detected merged cells, use them
        return table_data.get('merged_cells', [])
    
    def apply_merged_cells(self, table: Table, merged_cells: List[Dict]):
        """
        Apply merged cells to a DOCX table.
        
        Args:
            table: Table object
            merged_cells: List of merged cell definitions
        """
        for merge_def in merged_cells:
            row = merge_def['row']
            col = merge_def['col']
            row_span = merge_def.get('row_span', 1)
            col_span = merge_def.get('col_span', 1)
            
            if row < len(table.rows) and col < len(table.rows[row].cells):
                # Get the cell to merge from
                cell = table.rows[row].cells[col]
                
                # Merge cells
                if row_span > 1 or col_span > 1:
                    # Merge horizontally first
                    for c in range(1, col_span):
                        if col + c < len(table.rows[row].cells):
                            cell.merge(table.rows[row].cells[col + c])
                    
                    # Then merge vertically
                    for r in range(1, row_span):
                        if row + r < len(table.rows):
                            target_cell = table.rows[row + r].cells[col]
                            cell.merge(target_cell)
    
    def process_table_from_ocr(self, ocr_table: Dict, doc: Document) -> Table:
        """
        Complete table processing pipeline from OCR result.
        
        Args:
            ocr_table: Table data from OCR engine
            doc: Document object
            
        Returns:
            Created Table object
        """
        # Reconstruct table
        table = self.reconstruct_table(ocr_table, doc)
        
        # Detect and apply merged cells
        merged_cells = self.detect_merged_cells(ocr_table)
        if merged_cells:
            self.apply_merged_cells(table, merged_cells)
        
        return table

